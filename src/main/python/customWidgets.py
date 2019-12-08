from PyQt5.QtCore import Qt
from PyQt5 import QtWidgets, QtGui, uic, QtCore
import csv
from PyQt5.QtWidgets import QFileDialog
from pathlib import Path
import docx
import os 

from sqlitedb import DB
from lib.constants import *
from lib.database import VariableManager, QInterface


try:
    I=0
    CSV_DIALOG, _ = uic.loadUiType(BASEPATHS[I]+"ui/dialogs/importCsv.ui")
    MODALIDADES_DIALOD, _ = uic.loadUiType(BASEPATHS[I]+"ui/dialogs/modalidades.ui")
    NEW_MODALIDADE_WIDGET, _ = uic.loadUiType(BASEPATHS[I]+"ui/widgets/modalidadeForm.ui")
except:
    I=1
    CSV_DIALOG, _ = uic.loadUiType(BASEPATHS[I]+"ui/dialogs/importCsv.ui")
    MODALIDADES_DIALOD, _ = uic.loadUiType(BASEPATHS[I]+"ui/dialogs/modalidades.ui")
    NEW_MODALIDADE_WIDGET, _ = uic.loadUiType(BASEPATHS[I]+"ui/widgets/modalidadeForm.ui")

delimiter = CSV_SEPARATOR
SEPARADOR_CSV=CSV_SEPARATOR
types_of_encoding = ["utf-8", "cp1252"]

class Runner(QtCore.QThread):    
    def __init__(self, target, *args, **kwargs):
        super().__init__()
        self._target = target
        self._args = args
        self._kwargs = kwargs

    def run(self):
        if len(self._kwargs.keys())==0 and len(self._args)==0:
            self._target()   
        else:
            self._target(*self._args, **self._kwargs)



def nogui(func):
    from functools import wraps
    @wraps(func)
    def async_func(*args, **kwargs):
        runner = Runner(func, *args, **kwargs)
        # Keep the runner somewhere or it will be destroyed
        func.__runner = runner
        runner.start()

    return async_func
 


def messageDialog(iface=None, title="Concluído", info="", message=""):
    msgBox = QtWidgets.QMessageBox(iface)
    msgBox.setIcon(QtWidgets.QMessageBox.Question)
    msgBox.setWindowTitle(title)
    msgBox.setText(message)
    msgBox.setInformativeText(info)
    msgBox.setStandardButtons(QtWidgets.QMessageBox.Ok)
    msgBox.setDefaultButton(QtWidgets.QMessageBox.Ok)
   # msgBox.show()
    return msgBox.exec_() == QtWidgets.QMessageBox.Ok


def yesNoDialog(iface=None, title="Atenção", info="", message=""):
    msgBox = QtWidgets.QMessageBox(iface)
    msgBox.setIcon(QtWidgets.QMessageBox.Question)
    msgBox.setWindowTitle(title)
    msgBox.setText(message)
    msgBox.setInformativeText(info)
    msgBox.setStandardButtons(QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No)
    msgBox.setDefaultButton(QtWidgets.QMessageBox.No)
    #msgBox.show()
    return msgBox.exec_() == QtWidgets.QMessageBox.Yes


def confPath():
    path=QtCore.QStandardPaths.standardLocations(QtCore.QStandardPaths.ConfigLocation)[0] / Path(NAME)
    path.mkdir(parents=True, exist_ok=True)  
    return path


def tmpPath(): 
    path = QtCore.QStandardPaths.standardLocations(QtCore.QStandardPaths.TempLocation)[0] / Path(NAME)
    path.mkdir(parents=True, exist_ok=True)  
    return path



class dropDown(QtWidgets.QWidget):
    selected=QtCore.pyqtSignal(list)

    def __init__(self, lista=[], parent=None, text="", flags=Qt.WindowFlags()):
        super().__init__(parent=parent, flags=flags)
        self.layout=QtWidgets.QGridLayout(self)
        self.setLayout(self.layout)
        self.toolbutton = QtWidgets.QToolButton(self)
        self.layout.addWidget(self.toolbutton)
        self.toolbutton.setText(text)
        self.toolmenu = QtWidgets.QMenu()
        self.actions=[]
        self.checks=[]
        checkBox = QtWidgets.QCheckBox("Selecionar Todos", self.toolmenu)
        checkableAction = QtWidgets.QWidgetAction(self.toolmenu)
        checkableAction.setDefaultWidget(checkBox)
        self.toolmenu.addAction(checkableAction)
        checkBox.stateChanged.connect(lambda s: [c.setCheckState(s) for c in self.checks])
        checkBox.stateChanged.connect(lambda: self.selected.emit(self.selectedIndexes()))
        self.todos=checkBox
        self.repopulate(lista)                
        self.toolbutton.setMenu(self.toolmenu)
        self.toolbutton.setPopupMode(QtWidgets.QToolButton.InstantPopup)  
       #  self.setMinimumHeight(10)
       #  self.setMinimumWidth(100)

    def repopulate(self, lista):
        for act in self.actions:
            self.toolmenu.removeAction(act)
        self.actions=[]
        self.checks=[]
     
        for i in lista:            
            checkBox = QtWidgets.QCheckBox(str(i), self.toolmenu)
            checkBox.setChecked(True)
            checkBox.stateChanged.connect(lambda: self.selected.emit(self.selectedIndexes()))
            checkableAction = QtWidgets.QWidgetAction(self.toolmenu)
            checkableAction.setDefaultWidget(checkBox)
            checkableAction.setChecked(True)
            self.toolmenu.addAction(checkableAction)
            self.checks.append(checkBox)
            self.actions.append(checkableAction)

    def selectedIndexes(self):
        return [i for i, c in enumerate(self.checks) if c.isChecked()]
        
    def selectedTexts(self):
        return [c.text() for  c in self.checks if c.isChecked()]
       

class csvDialog(QtWidgets.QDialog, CSV_DIALOG):
    def __init__(self, dataNamesList:list, parent=None, file=None):
        '''dataNamesList: lista de strings com os nomes dos possíveis atributos '''
        super().__init__(None)
        self.setupUi(self)
        self.horizontalLayout : QtWidgets.QHBoxLayout
        self.cbList=[]
        self.dataNamesList=dataNamesList         
        self.result=[]
        self.header=[]
        self.defaultFile=file

        for i,dName in enumerate(dataNamesList):
            cb=QtWidgets.QComboBox()                       
            cb.addItems(dataNamesList)
            self.con(cb,i) 
            vlay=QtWidgets.QVBoxLayout()
            lbl=QtWidgets.QLabel(str(i+1))
            vlay.addWidget(lbl)
            vlay.addWidget(cb)
            self.horizontalLayout.addLayout(vlay)
            cb.show()            
            lbl.show()
            self.cbList.append(cb)           

    def updateData(self, text, index):        
        for i, cb in enumerate(self.cbList):
            cb:QtWidgets.QComboBox         
            if i!=index:
                j=cb.currentIndex()
                while self.dataNamesList[j] in [cb.currentText() for k, cb in enumerate(self.cbList) if i!=k]:
                    j=j+1 if not j+1>=len(self.dataNamesList) else 0
                cb.setCurrentIndex(j)                
                self.con(cb, i) 
    
    def default(self):
        range=QtWidgets.QTableWidgetSelectionRange(0,0,9,len(self.dataNamesList)-1)
        self.tableWidget.setRangeSelected(range, True)
        for i,cb in enumerate(self.cbList):                        
            cb : QtWidgets.QComboBox
            cb.currentTextChanged.disconnect()
            cb.setCurrentIndex(i)
            self.con(cb,i) 

    def con(self, cb, i):
        cb.currentTextChanged.connect(lambda text: self.updateData(text, i))       

    def exec_(self):
        if self.openFile():
            self.default()
            return super().exec_()
        else:
            messageDialog(title="Atenção!", message="A planilha não contém um número de colunas suficiente ("+str(len(self.dataNamesList))+")")
            return                   
 

    def show(self):
        if self.openFile():
            self.default()
            return super().show()
        else:
            messageDialog(title="Atenção!", message="A planilha não contém um número de colunas suficiente ("+str(len(self.dataNamesList))+")")
            return        
    
    def accept(self):
        self.result=self.csvRead()   
        if not self.result:
            return 
        return super().accept()
    
    def csvRead(self):
        global delimiter
        result=[]
        first=True
        columnIndexes=list(set(index.column() for index in self.tableWidget.selectionModel().selectedIndexes()))
        if len(columnIndexes) != len(self.dataNamesList):
            messageDialog(title="Atenção!", message="Por favor selectione o mesmo número de colunas que de campos necessários \n Você seleciou "+str(len(columnIndexes))+", mas são necessários "+str(len(self.dataNamesList)))
            return False
        import codecs
        for encoding_type in types_of_encoding: 
            try:
                result=[]
                with codecs.open(self.filepath, 'r',  encoding= encoding_type, errors ='replace') as fi:
                    for i, r in enumerate(csv.reader(fi, delimiter=delimiter)):                        
                        if self.checkBox.isChecked() and first:
                            first=False
                            continue
                        first=False
                        dado = {}
                        j=0
                        for field, dName in zip([f for index, f in enumerate(r) if index in columnIndexes], 
                                                                [cb.currentText() for cb in self.cbList]):
                            field=str(field)
                            if dName=="":
                                messageDialog(title="Atenção!", message="Por favor atribua valores para cada coluna")
                                return False
                            dado[dName]=u"%s"%field
                        result.append(dado)
                break
            except:
                continue

        return result  

    def openFile(self):
        filename = QtWidgets.QFileDialog.getOpenFileName(filter="Arquivo/Planilha CSV (*.csv)")[0] if self.defaultFile is None else self.defaultFile
        if filename in ["", None]: return False
        self.filepath=filename
        self.tableWidget : QtWidgets.QTableWidget
        import codecs
        for encoding_type in types_of_encoding:    
            try:
                with codecs.open(self.filepath, 'r', encoding= encoding_type, errors ='replace') as fi:
                    for i,r in enumerate(csv.reader(fi, delimiter=delimiter)):        
                        for j,field in enumerate(r):
                            field=str(field)
                            if i==0:
                                self.header.append(field)                        
                                continue
                            if i>=1:
                                if i==1 and j==0:
                                    self.tableWidget.setRowCount(10)                      
                                    self.tableWidget.setColumnCount(len(self.header))
                                    self.tableWidget.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectColumns)
                                    self.tableWidget.setSelectionMode(QtWidgets.QAbstractItemView.MultiSelection)
                                    #self.tableWidget.setHorizontalHeaderLabels((str(i+1) for i in range(len(self.header)) ))
                                    self.tableWidget.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Interactive)
                                    self.tableWidget.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                                    self.tableWidget.horizontalHeader().setStretchLastSection(True)                          

                                    for k,f in enumerate(self.header):        
                                        tableItem=QtWidgets.QTableWidgetItem(u"%s" % str(f))
                                        tableItem.setFlags(tableItem.flags() ^ Qt.ItemIsEditable)
                                        self.tableWidget.setItem(0,k,tableItem)

                                tableItem=QtWidgets.QTableWidgetItem(u"%s" % str(u"%s"%field))
                                tableItem.setFlags(tableItem.flags() ^ Qt.ItemIsEditable)
                                self.tableWidget.setItem(i,j,tableItem)
                        if i>10:
                            break
                break
            except:
                continue
        self.stretchTable(self.tableWidget)
        return len(self.dataNamesList)<=len(self.header)

    def stretchTable(self, table):            
            tableSize = table.width()
            sideHeaderWidth = table.verticalHeader().width()
            tableSize -= sideHeaderWidth
            numberOfColumns = table.columnCount()

            remainingWidth = tableSize % numberOfColumns
            for columnNum in range(table.columnCount()):
                if remainingWidth > 0:
                    table.setColumnWidth(columnNum, int(tableSize / numberOfColumns) + 1)
                    remainingWidth -= 1
                else:
                    table.setColumnWidth(columnNum, int(tableSize / numberOfColumns))

    def resizeEvent(self, event):
        self.stretchTable(self.tableWidget)
        super().resizeEvent(event)  # Restores the original behaviour of the resize event


def exportCsv(listaDeAlunos, countChanged, finished):
    '''
    listaDeAlunos: lista de Dicionários de alunos
    filename: Caminho do arquivo csv a salvar
    '''
    filename=QFileDialog.getSaveFileName(filter="Arquivo CSV (*.csv)")[0]
    #filename=["/home/matheus/test.csv"]
    if not filename: return
    filename = filename if filename.endswith(".csv") else filename+".csv"
    header=list(listaDeAlunos[0].keys())    
    with open(filename, "w", encoding=types_of_encoding[1]) as fo:
        writer = csv.writer(fo, delimiter=CSV_SEPARATOR, dialect='excel')
        if type(header)==list:
            writer.writerow(header)
        for i,dic in enumerate(listaDeAlunos):
            countChanged.emit(int(i/len(listaDeAlunos)*100),"Exportando csv")
            r=list(dic.values())
            writer.writerow(r)
    
    if yesNoDialog(message="Criar tabela no word?"):
        exportDoc(filename, countChanged, finished)
    else:
        finished.emit()
    
@nogui  
def exportDoc(filename, countChanged, finished, k=None):
    try:
        I=0
        doc = docx.Document(BASEPATHS[I]+"templates/default.docx")
    except:
        I=1
        doc = docx.Document(BASEPATHS[I]+"templates/default.docx")

    with open(filename, "r", newline='',encoding=types_of_encoding[1]) as f:
        indexes = [0,2,5,6,7,8,9]
        csv_headers=["Nome", "Nascimento", "Mãe", "Pai", "Telefone", "Endereço", "Turma"]
        csv_reader = csv.reader(f, delimiter=CSV_SEPARATOR)       
        csv_cols = len(csv_headers)
        table = doc.add_table(rows=2, cols=csv_cols)
        hdr_cells = table.rows[0].cells
        for i in range(csv_cols):           
            hdr_cells[i].text = csv_headers[i]
        row_count = 1000
        csv_reader = csv.reader(f, delimiter=CSV_SEPARATOR)       
        headers=len(next(csv_reader))
        for ci,row in enumerate(csv_reader):
            countChanged.emit(int(ci/row_count*80),"Exportando docx")
            if len(row) != headers:
                continue
            row_cells = table.add_row().cells
            for i in range(headers):
                if i in indexes:
                    row_cells[indexes.index(i)].text = row[i]

        paragraph =hdr_cells[0].paragraphs[0]
        run = paragraph.runs
        for ci,row in enumerate(table.rows):
            countChanged.emit(int(ci/row_count*20)+80,"Melhorando Layout")
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font                   
                        font.name = 'Times New Roman'
                        font.size= docx.shared.Pt(10)

    doc.add_page_break()        
    doc.save(filename[:-4]+".docx")
    finished.emit()


class NewModalidadeWidget(QtWidgets.QWidget, NEW_MODALIDADE_WIDGET):
    def __init__(self, iface):
        QtWidgets.QWidget.__init__(self)
        NEW_MODALIDADE_WIDGET.__init__(self)
        self.setupUi(self)        
        self.lineEdit:QtWidgets.QLineEdit
        self.lineEdit.setPlaceholderText("Nome da Modalidade")

class NewTurmaWidget(QtWidgets.QWidget, NEW_MODALIDADE_WIDGET):
    def __init__(self, iface):
        QtWidgets.QWidget.__init__(self)
        NEW_MODALIDADE_WIDGET.__init__(self)
        self.setupUi(self)
        self.lineEdit:QtWidgets.QLineEdit
        self.lineEdit.setPlaceholderText("Turma/Série")

class ModalidadesDialog(QtWidgets.QDialog, MODALIDADES_DIALOD):
    edited = QtCore.pyqtSignal()
    def __init__(self, iface):
        QtWidgets.QDialog.__init__(self)
        MODALIDADES_DIALOD.__init__(self)
        self.setupUi(self)

        self.iface:MainWindow
        self.varManager:VariableManager
        self.listWidget:QtWidgets.QListWidget
        self.listWidget_2:QtWidgets.QListWidget
        self.modalidades:QInterface
        self.buttonBox:QtWidgets.QDialogButtonBox

        self.iface=iface
        self.varManager:VariableManager = iface.varManager
        self.modalidades=self.iface.modalidades    
        self.listWidget.itemClicked.connect(self.modalidadeChanged)

        self.modalidades.setup(self,
            signals =  [],
            slots   =  [], 
            properties=["modalidades"], 
            readers  = [self.read],
            writers  = [self.write]
        )

    def read(self):
        return []

    def write(self,modalidades:list):
        for i,m in enumerate(modalidades):   
            self.addToListWidget1(m, removable=True if i<len(modalidades)-1 else False)
        
    def addToListWidget1(self, modalidade, removable=False):

        itemN = QtWidgets.QListWidgetItem() 
        widget = NewModalidadeWidget(self)
        widget.label.setText("Modalidade Escolar: ")
        widget.horizontalLayout.addStretch()
        widget.horizontalLayout.setSizeConstraint(QtWidgets.QLayout.SetFixedSize)
        itemN.setSizeHint(widget.sizeHint())  
        self.listWidget.addItem(itemN)
        self.listWidget.setItemWidget(itemN, widget)

        i=self.listWidget.row(itemN)
        widget.lineEdit.setText(modalidade.nome)                
        widget.pushButton.clicked.connect(lambda: self.addToListWidget1(Modalidade()))
        widget.pushButton.clicked.connect(lambda: self.setRemovableFromListWidget1(widget, itemN))

        if removable:
            self.setRemovableFromListWidget1(widget, itemN)


    def removeFromListWidget1(self, item):
        i=self.listWidget.row(item)
        self.listWidget.takeItem(i)
        self.modalidades.get().removeByIndex(i)

    def setRemovableFromListWidget1(self, widget, itemN):
        while True:
            try: widget.pushButton.clicked.disconnect() 
            except Exception: break
        widget.lineEdit.setReadOnly(True)
        widget.pushButton.clicked.connect(lambda: self.removeFromListWidget1(itemN))
        widget.pushButton.setText("Remover")

    def modalidadeChanged(self):
        item=self.listWidget.currentItem()
        self.listWidget_2.clear()  
        lista=self.modalidades.get().modalidades[self.listWidget.row(item)].turmas
        for i,t in enumerate(lista):
            self.addToListWidget2(t,removable=True if i < len(lista)-1 else False)

    def addToListWidget2(self, turma, removable=False):        
        item=self.listWidget.currentItem()
        itemN = QtWidgets.QListWidgetItem() 
        widget = NewTurmaWidget(self)
        widget.label.setText("Turma: ")
        widget.horizontalLayout.addStretch()
        widget.horizontalLayout.setSizeConstraint(QtWidgets.QLayout.SetFixedSize)
        itemN.setSizeHint(widget.sizeHint())    

        self.listWidget_2.addItem(itemN)
        self.listWidget_2.setItemWidget(itemN, widget)
        widget.pushButton.clicked.connect(lambda: self.addToListWidget2(Turma()))
        widget.pushButton.clicked.connect(lambda: self.setRemovableFromListWidget2(widget, itemN))

        if removable:
            self.setRemovableFromListWidget2(widget, itemN)

    def removeFromListWidget2(self, item):
        i=self.listWidget_2.row(item)
        item=self.listWidget.currentItem()       
        self.listWidget_2.takeItem(i)   
        self.modalidades.get().modalidades[self.listWidget.row(item)]
        

    def setRemovableFromListWidget2(self, widget, itemN):
        while True:
            try: widget.pushButton.clicked.disconnect() 
            except Exception: break
        widget.lineEdit.setReadOnly(True)
        widget.pushButton.clicked.connect(lambda: self.removeFromListWidget2(itemN))
        widget.pushButton.setText("Remover")


def test1(*args):
    app = QtWidgets.QApplication(*args)
    app.restart=False
    d1={"nome":'majose', "matricula":"ER215", "dataNasc":'12/05/87', "RG":'askfasj1545', "CPF":'15618684',
      "nomeDaMae":'josefina', "nomeDoPai":'Jão', "telefone":'121839128', "endereco":'fksdkf 239j 29r',
 "serie":2, "escola":1, "idade":13, "lat":-19.231, "long":47.12331}
    d2={"nome":'matheus', "matricula":"ER128", "dataNasc":'17/05/87', "RG":'askfasj1545', "CPF":'15618684',
      "nomeDaMae":'josefina', "nomeDoPai":'Jão', "telefone":'121839128', "endereco":'fksdkf 239j 29r',
 "serie":5, "escola":1, "idade":21, "lat":-19.231, "long":47.12331}
    d3={"nome":'carlos', "matricula":"ER125", "dataNasc":'18/05/87', "RG":'askfasj1545', "CPF":'15618684',
      "nomeDaMae":'josefina', "nomeDoPai":'Jão', "telefone":'121839128', "endereco":'fksdkf 239j 29r',
 "serie":8, "escola":3, "idade":15, "lat":-19.231, "long":47.12331}
 
    exportCsv([d1,d2,d3])
    r=app.exec_()        
    return r


def test(*args):  
        app = QtWidgets.QApplication(*args)
        app.restart=False
        win=csvDialog(["nome", "idade", "x"], app)
        win.show()     
        r=app.exec_()        
        print(win.result)
        return r



if __name__ == '__main__':
    import sys
    currentExitCode=test1(sys.argv)
    sys.exit(currentExitCode)
